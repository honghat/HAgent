"""PDF tools — create, edit pages, from images/word, translate."""

from __future__ import annotations

import asyncio
import base64
import io
import json
import re
import tempfile
import time
import uuid
from pathlib import Path
from urllib.parse import quote

import aiohttp
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from PIL import Image, ImageEnhance, ImageOps
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

try:
    from sse_starlette.sse import EventSourceResponse
except ImportError:  # fallback nếu thiếu sse-starlette
    EventSourceResponse = None

router = APIRouter(prefix="/api/pdf", tags=["pdf"])


@router.get("/font/noto-sans")
async def pdf_unicode_font():
    """Return a TTF font that pdf-lib can embed for Vietnamese text edits."""
    font_path = Path(__file__).resolve().parent.parent.parent / "assets" / "fonts" / "DejaVuSans.ttf"
    if not font_path.exists():
        raise HTTPException(404, "Không tìm thấy font Unicode cho PDF editor")
    return FileResponse(
        str(font_path),
        media_type="font/ttf",
        filename="DejaVuSans.ttf",
    )


# === LLM correction layer (multi-provider) ===
# Cấu hình LLM do user chọn trong settings UI của PdfEditor; truyền qua
# request body khi gọi translate. Không lưu key trên server (per-user).
#
# Providers: deepseek | gemini | openai | ollama | lmstudio | custom | none
# - deepseek/openai/gemini: dùng OpenAI-compatible hoặc Gemini REST.
# - ollama: http://localhost:11434/v1 (OpenAI-compatible).
# - lmstudio: http://localhost:1234/v1.
# - custom: base_url + key tuỳ ý (9router, proxy khác).
# - none: skip LLM, dùng lingva.ml cũ (giữ tương thích ngược).


async def _llm_chat(messages: list[dict], settings: dict, timeout: int = 60) -> str:
    """Gọi LLM completion trả về text. Raises HTTPException nếu lỗi."""
    provider = (settings.get("provider") or "none").lower()
    model = (settings.get("model") or "").strip()
    api_key = (settings.get("api_key") or "").strip()
    base_url = (settings.get("base_url") or "").strip().rstrip("/")

    if provider in ("none", "", "off"):
        raise HTTPException(400, "LLM chưa bật (provider=none)")
    if not model:
        raise HTTPException(400, "Thiếu model name")

    headers = {"Content-Type": "application/json"}
    payload: dict
    url: str

    if provider == "gemini":
        # Gemini dùng REST riêng: .../v1beta/models/{model}:generateContent?key=KEY
        if not api_key:
            raise HTTPException(400, "Gemini cần api_key")
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        # Convert OpenAI messages sang Gemini contents
        contents = []
        for m in messages:
            role = "user" if m["role"] in ("user", "system") else "model"
            if role == "user" and contents and contents[-1]["role"] == "user":
                contents[-1]["parts"].append({"text": m["content"]})
            else:
                contents.append({"role": role, "parts": [{"text": m["content"]}]})
        if not contents or contents[0]["role"] != "user":
            contents.insert(0, {"role": "user", "parts": ["OK"]})
        payload = {"contents": contents, "generationConfig": {"temperature": 0.2}}
    else:
        # OpenAI-compatible (deepseek/openai/ollama/lmstudio/custom)
        if not base_url:
            defaults = {
                "deepseek": "https://api.deepseek.com/v1",
                "openai": "https://api.openai.com/v1",
                "ollama": "http://localhost:11434/v1",
                "lmstudio": "http://localhost:1234/v1",
            }
            base_url = defaults.get(provider, "")
        if not base_url:
            raise HTTPException(400, f"Provider {provider} cần base_url")
        url = f"{base_url}/chat/completions"
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": messages,
            "temperature": 0.2,
            "max_tokens": 4096,
        }
        if provider == "ollama":
            payload.pop("max_tokens", None)

    try:
        async with aiohttp.ClientSession() as s:
            async with s.post(
                url, headers=headers, json=payload,
                timeout=aiohttp.ClientTimeout(total=timeout),
            ) as r:
                txt = await r.text()
                if r.status >= 400:
                    raise HTTPException(
                        r.status,
                        f"LLM {provider} lỗi {r.status}: {txt[:200]}",
                    )
                data = json.loads(txt)
                if provider == "gemini":
                    candidates = data.get("candidates") or []
                    if not candidates:
                        raise HTTPException(502, "Gemini trả rỗng")
                    parts = candidates[0].get("content", {}).get("parts", [])
                    out = "".join(p.get("text", "") for p in parts)
                else:
                    choices = data.get("choices") or []
                    if not choices:
                        raise HTTPException(502, f"{provider} trả rỗng")
                    out = choices[0].get("message", {}).get("content", "") or ""
                return out.strip()
    except HTTPException:
        raise
    except asyncio.TimeoutError:
        raise HTTPException(504, f"LLM {provider} timeout {timeout}s")
    except aiohttp.ClientError as exc:
        raise HTTPException(502, f"LLM {provider} kết nối lỗi: {exc}")


def _split_into_chunks(text: str, max_chars: int = 1500) -> list[str]:
    """Chia text theo đoạn văn/dòng, mỗi chunk <= max_chars. Tránh cắt giữa từ."""
    if not text.strip():
        return []
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    buf = ""
    # Ưu tiên tách theo đoạn (\n\n), rồi câu (. ! ? \n), rồi từ
    paragraphs = re.split(r"(\n\s*\n)", text)
    for para in paragraphs:
        if not para:
            continue
        if len(buf) + len(para) <= max_chars:
            buf += para
        else:
            if buf.strip():
                chunks.append(buf.strip())
                buf = ""
            if len(para) > max_chars:
                # Cắt câu dài
                sentences = re.split(r"(?<=[.!?\n])\s+", para)
                for sent in sentences:
                    if len(buf) + len(sent) <= max_chars:
                        buf += sent + " "
                    else:
                        if buf.strip():
                            chunks.append(buf.strip())
                            buf = ""
                        if len(sent) > max_chars:
                            # Cắt từ cuối
                            for i in range(0, len(sent), max_chars):
                                chunks.append(sent[i:i + max_chars].strip())
                        else:
                            buf = sent + " "
            else:
                buf = para
    if buf.strip():
        chunks.append(buf.strip())
    return [c for c in chunks if c]


async def _llm_correct_text(raw_text: str, settings: dict) -> str:
    """Làm sạch OCR text: sửa typo tiếng Việt, tách đoạn dính, giữ nguyên ngôn ngữ gốc."""
    if not raw_text.strip():
        return raw_text
    system = (
        "Bạn là biên tập viên tiếng Việt chuyên sửa văn bản OCR.\n"
        "Quy tắc:\n"
        "1. Sửa lỗi chính tả, lỗi dấu thanh, lỗi đầu cuối từ do OCR gây ra.\n"
        "2. Tách các đoạn/dòng bị dính vào nhau thành các đoạn riêng biệt.\n"
        "3. Giữ nguyên ngôn ngữ gốc của văn bản (không dịch).\n"
        "4. Giữ nguyên nội dung, ý nghĩa, số liệu, tên riêng.\n"
        "5. Trả về CHỈ văn bản đã sửa, KHÔNG giải thích, KHÔNG tiêu đề thêm."
    )
    chunks = _split_into_chunks(raw_text, max_chars=2500)
    out_parts: list[str] = []
    for chunk in chunks:
        try:
            corrected = await _llm_chat(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": chunk},
                ],
                settings,
                timeout=90,
            )
            out_parts.append(corrected or chunk)
        except HTTPException:
            # Nếu LLM fail 1 chunk, giữ nguyên chunk đó
            out_parts.append(chunk)
    return "\n\n".join(out_parts)


async def _llm_translate_chunk(text: str, target_lang: str, settings: dict) -> str:
    """Dịch text qua LLM, fallback lingva.ml nếu LLM fail/disabled."""
    if not text.strip():
        return text
    provider = (settings.get("provider") or "none").lower()
    if provider in ("none", "", "off"):
        return await _translate_chunk(text, target_lang)
    lang_name = {
        "vi": "tiếng Việt", "en": "tiếng Anh", "ja": "tiếng Nhật",
        "ko": "tiếng Hàn", "zh": "tiếng Trung", "fr": "tiếng Pháp",
        "de": "tiếng Đức", "es": "tiếng Tây Ban Nha", "ru": "tiếng Nga",
    }.get(target_lang.lower(), target_lang)
    system = (
        f"Bạn là dịch giả chuyên nghiệp. Dịch văn bản sang {lang_name}.\n"
        "Quy tắc:\n"
        "1. Dịch tự nhiên, trôi chảy, giữ đúng thuật ngữ chuyên ngành.\n"
        "2. Giữ nguyên format: tiêu đề, danh sách, xuống dòng.\n"
        "3. Không thêm giải thích, không bình luận.\n"
        "4. Chỉ trả về bản dịch."
    )
    try:
        return await _llm_chat(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": text},
            ],
            settings,
            timeout=90,
        )
    except HTTPException:
        return await _translate_chunk(text, target_lang)

# In-memory job store cho translate-job (SSE progress)
_translate_jobs: dict[str, dict] = {}
_translate_jobs_lock = asyncio.Lock()
_JOB_TTL = 600  # 10 phút

LANG_MAP = {
    "vi": "vi", "en": "en", "ja": "ja", "ko": "ko", "zh": "zh",
    "fr": "fr", "de": "de", "es": "es", "ru": "ru", "th": "th",
}


def _check_size(file_bytes: bytes) -> None:
    return None


def _pdf_response(buf: bytes, filename: str = "output.pdf") -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(buf),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{quote(filename)}"'},
    )


def _pdf_data_dir() -> Path:
    project_root = Path(__file__).resolve().parent.parent.parent.parent
    d = project_root / "data" / "pdf"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _save_to_recent(buf: bytes, filename: str) -> str:
    """Lưu PDF vào data/pdf/ để xuất hiện trong 'Gần đây'. Trả về absolute path."""
    safe = re.sub(r"[^\w.\- ]+", "_", filename).strip() or "output.pdf"
    if not safe.lower().endswith(".pdf"):
        safe += ".pdf"
    target = _pdf_data_dir() / safe
    if target.exists():
        stem, suf = target.stem, target.suffix
        i = 2
        while True:
            candidate = target.with_name(f"{stem} ({i}){suf}")
            if not candidate.exists():
                target = candidate
                break
            i += 1
    target.write_bytes(buf)
    return str(target.resolve())


def _pdf_response_saved(buf: bytes, filename: str = "output.pdf") -> StreamingResponse:
    saved_path = _save_to_recent(buf, filename)
    return StreamingResponse(
        io.BytesIO(buf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{quote(filename)}"; filename*=UTF-8\'\'{quote(filename)}"',
            "X-Saved-Path": saved_path,
        },
    )


_FONT_REG = None


def _register_fonts() -> str:
    """Đăng ký font hỗ trợ Unicode (kể cả tiếng Việt có dấu). Trả về family name.

    Tìm theo thứ tự:
      1. Liberation Sans trong assets/fonts/ (cross-platform, nhưng KHÔNG có tiếng Việt)
      2. DejaVu Sans trong assets/fonts/ (cross-platform, CÓ tiếng Việt - cần download)
      3. Verdana/Tahoma/Arial Unicode trên macOS
      4. Fallback Helvetica (sẽ mất dấu tiếng Việt)
    """
    global _FONT_REG
    if _FONT_REG is not None:
        return _FONT_REG
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.pdfmetrics import registerFontFamily

        here = Path(__file__).resolve().parent.parent.parent
        fonts_dir = here / "assets" / "fonts"

        candidates: list[tuple[str, str, str, str]] = [
            # (stem, normal_filename, bold_filename, italic_filename)
            ("LibSans", "LiberationSans-Regular.ttf", "LiberationSans-Bold.ttf", "LiberationSans-Italic.ttf"),
            ("DejaVu", "DejaVuSans.ttf", "DejaVuSans-Bold.ttf", "DejaVuSans-Oblique.ttf"),
        ]
        for stem, reg_f, bold_f, ital_f in candidates:
            if (fonts_dir / reg_f).exists() and (fonts_dir / bold_f).exists():
                pdfmetrics.registerFont(TTFont(stem, str(fonts_dir / reg_f)))
                pdfmetrics.registerFont(TTFont(f"{stem}-Bold", str(fonts_dir / bold_f)))
                if (fonts_dir / ital_f).exists():
                    pdfmetrics.registerFont(TTFont(f"{stem}-Italic", str(fonts_dir / ital_f)))
                try:
                    registerFontFamily(stem, normal=stem, bold=f"{stem}-Bold",
                                       italic=f"{stem}-Italic" if (fonts_dir / ital_f).exists() else stem,
                                       boldItalic=f"{stem}-Bold")
                except Exception:
                    pass
                _FONT_REG = stem
                return _FONT_REG

        # Fallback: macOS system fonts (hỗ trợ tiếng Việt)
        mac_paths = [
            ("MacVerdana", "/System/Library/Fonts/Supplemental/Verdana.ttf",
             "/System/Library/Fonts/Supplemental/Verdana Bold.ttf"),
            ("MacTahoma", "/System/Library/Fonts/Supplemental/Tahoma.ttf",
             "/System/Library/Fonts/Supplemental/Tahoma Bold.ttf"),
        ]
        for stem, reg_p, bold_p in mac_paths:
            if Path(reg_p).exists() and Path(bold_p).exists():
                pdfmetrics.registerFont(TTFont(stem, reg_p))
                pdfmetrics.registerFont(TTFont(f"{stem}-Bold", bold_p))
                try:
                    registerFontFamily(stem, normal=stem, bold=f"{stem}-Bold",
                                       italic=stem, boldItalic=f"{stem}-Bold")
                except Exception:
                    pass
                _FONT_REG = stem
                return _FONT_REG

        # Fallback Linux: DejaVu Sans toàn hệ thống
        linux_paths = [
            ("SysDejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
             "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ]
        for stem, reg_p, bold_p in linux_paths:
            if Path(reg_p).exists() and Path(bold_p).exists():
                pdfmetrics.registerFont(TTFont(stem, reg_p))
                pdfmetrics.registerFont(TTFont(f"{stem}-Bold", bold_p))
                try:
                    registerFontFamily(stem, normal=stem, bold=f"{stem}-Bold",
                                       italic=stem, boldItalic=f"{stem}-Bold")
                except Exception:
                    pass
                _FONT_REG = stem
                return _FONT_REG

        _FONT_REG = "Helvetica"
    except Exception as e:
        import sys
        print(f"[pdf_tools] font registration failed: {e}", file=sys.stderr)
        _FONT_REG = "Helvetica"
    return _FONT_REG


def _build_text_pdf(text: str, *, title: str | None = None) -> bytes:
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, topMargin=40, bottomMargin=40, leftMargin=40, rightMargin=40)
    styles = getSampleStyleSheet()
    fname = _register_fonts()
    body = styles["BodyText"]
    body.fontName = fname
    body.fontSize = 11
    body.leading = 16
    story = []
    if title:
        h = styles["Heading2"]
        h.fontName = "LibSans-Bold" if fname == "LibSans" else fname
        story.append(Paragraph(title, h))
        story.append(Spacer(1, 12))
    for para in text.split("\n\n"):
        clean = para.replace("\n", "<br/>")
        clean = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", clean)
        story.append(Paragraph(clean or "&nbsp;", body))
        story.append(Spacer(1, 8))
    doc.build(story)
    return out.getvalue()


@router.post("/blank")
async def blank_pdf(pages: int = Form(1)):
    n = max(1, min(int(pages or 1), 50))
    writer = PdfWriter()
    for _ in range(n):
        writer.add_blank_page(width=595, height=842)  # A4 in pts
    out = io.BytesIO()
    writer.write(out)
    return _pdf_response_saved(out.getvalue(), "blank.pdf")


@router.post("/from-text")
async def pdf_from_text(text: str = Form(...), title: str = Form("")):
    if not text.strip():
        raise HTTPException(400, "Text rỗng")
    buf = _build_text_pdf(text, title=title or None)
    fname = (title.strip() or "from-text") + ".pdf"
    return _pdf_response_saved(buf, fname)


@router.post("/from-images")
async def pdf_from_images(files: list[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "Chưa chọn ảnh")
    images = []
    for f in files:
        raw = await f.read()
        _check_size(raw)
        try:
            img = Image.open(io.BytesIO(raw)).convert("RGB")
        except Exception as exc:
            raise HTTPException(400, f"Ảnh không hợp lệ: {f.filename}") from exc
        images.append(img)
    if not images:
        raise HTTPException(400, "Không đọc được ảnh nào")
    out = io.BytesIO()
    first, rest = images[0], images[1:]
    first.save(out, format="PDF", save_all=True, append_images=rest)
    fname = Path(files[0].filename or "images").stem + ".pdf"
    return _pdf_response_saved(out.getvalue(), fname)


@router.post("/from-docx")
async def pdf_from_docx(file: UploadFile = File(...)):
    raw = await file.read()
    _check_size(raw)
    if not file.filename or not file.filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(400, "Chỉ nhận file .docx")
    from docx import Document
    try:
        doc = Document(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"Không đọc được docx: {exc}") from exc
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text:
        raise HTTPException(400, "File rỗng hoặc không lấy được text")
    out_name = Path(file.filename).stem + ".pdf"
    return _pdf_response_saved(_build_text_pdf(text), out_name)


@router.post("/reorder")
async def reorder_pages(file: UploadFile = File(...), order: str = Form(...)):
    """Rewrite PDF keeping only pages in `order` (1-based, comma list)."""
    raw = await file.read()
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc
    total = len(reader.pages)
    indices = []
    for tok in order.split(','):
        tok = tok.strip()
        if not tok:
            continue
        try:
            n = int(tok)
        except ValueError:
            continue
        if 1 <= n <= total:
            indices.append(n - 1)
    if not indices:
        raise HTTPException(400, 'Thứ tự trang rỗng/không hợp lệ')
    rotations = {}
    # Optional rotation map encoded as "1:90,3:180" appended after a pipe
    writer = PdfWriter()
    for idx in indices:
        page = reader.pages[idx]
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return _pdf_response(out.getvalue(), 'edited.pdf')


def _parse_page_targets(pages: str, total: int) -> set[int]:
    """Parse danh sách trang 1-based, hỗ trợ "1,3-5"; rỗng nghĩa là tất cả."""
    targets: set[int] = set()
    for raw_part in (pages or "").split(","):
        part = raw_part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                start, end = [int(x.strip()) for x in part.split("-", 1)]
            except ValueError:
                continue
            if start > end:
                start, end = end, start
            for page_no in range(start, end + 1):
                if 1 <= page_no <= total:
                    targets.add(page_no - 1)
            continue
        try:
            page_no = int(part)
        except ValueError:
            continue
        if 1 <= page_no <= total:
            targets.add(page_no - 1)
    return targets or set(range(total))


@router.post("/rotate-small")
async def rotate_small_pdf(
    file: UploadFile = File(...),
    angle: float = Form(...),
    pages: str = Form(""),
):
    """Xoay rất nhỏ bằng cách raster hóa trang, phù hợp PDF scan/ảnh."""
    raw = await file.read()
    _check_size(raw)
    if not raw.startswith(b"%PDF"):
        raise HTTPException(400, "File không phải PDF")
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc

    total = len(reader.pages)
    if total == 0:
        raise HTTPException(400, "PDF không có trang")
    rotate_angle = float(angle or 0)
    targets = _parse_page_targets(pages, total)

    writer = PdfWriter()
    try:
        resample = Image.Resampling.BICUBIC
    except AttributeError:
        resample = Image.BICUBIC

    for idx, src_page in enumerate(reader.pages):
        if idx not in targets or abs(rotate_angle) < 0.01:
            writer.add_page(src_page)
            continue

        png = _render_page_png(raw, idx, scale=1.0)
        if not png:
            raise HTTPException(500, f"Không render được trang {idx + 1}")

        img = Image.open(io.BytesIO(png)).convert("RGB")
        fill_color = _average_corner_color(img)
        img = img.rotate(
            -rotate_angle,
            resample=resample,
            expand=False,
            fillcolor=fill_color,
        )

        width = float(src_page.mediabox.width)
        height = float(src_page.mediabox.height)
        img_buf = io.BytesIO()
        img.save(img_buf, format="JPEG", quality=88)
        img_buf.seek(0)

        page_pdf = io.BytesIO()
        c = canvas.Canvas(page_pdf, pagesize=(width, height))
        c.drawImage(ImageReader(img_buf), 0, 0, width=width, height=height)
        c.showPage()
        c.save()
        rotated_reader = PdfReader(io.BytesIO(page_pdf.getvalue()))
        writer.add_page(rotated_reader.pages[0])

    out = io.BytesIO()
    writer.write(out)
    return _pdf_response(out.getvalue(), "rotated.pdf")


@router.post("/scan-page")
async def scan_page_pdf(
    file: UploadFile = File(...),
    page: int = Form(...),
    points: str = Form(...),
    enhance: bool = Form(True),
):
    """Nắn phối cảnh một trang theo 4 điểm normalized: tl,tr,br,bl."""
    raw = await file.read()
    _check_size(raw)
    if not raw.startswith(b"%PDF"):
        raise HTTPException(400, "File không phải PDF")
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc

    total = len(reader.pages)
    page_idx = int(page or 1) - 1
    if page_idx < 0 or page_idx >= total:
        raise HTTPException(400, "Trang không hợp lệ")

    try:
        raw_points = json.loads(points)
        if not isinstance(raw_points, list) or len(raw_points) != 4:
            raise ValueError("points cần 4 điểm")
        norm_points = [
            (
                max(0.0, min(1.0, float(p["x"]))),
                max(0.0, min(1.0, float(p["y"]))),
            )
            for p in raw_points
        ]
    except Exception as exc:
        raise HTTPException(400, f"points không hợp lệ: {exc}") from exc

    png = _render_page_png(raw, page_idx, scale=1.0)
    if not png:
        raise HTTPException(500, f"Không render được trang {page_idx + 1}")
    img = Image.open(io.BytesIO(png)).convert("RGB")
    img_w, img_h = img.size
    src_points = [(x * img_w, y * img_h) for x, y in norm_points]
    dst_points = [(0, 0), (img_w, 0), (img_w, img_h), (0, img_h)]
    coeffs = _perspective_coeffs(dst_points, src_points)
    try:
        perspective_mode = Image.Transform.PERSPECTIVE
        resample = Image.Resampling.BICUBIC
    except AttributeError:
        perspective_mode = Image.PERSPECTIVE
        resample = Image.BICUBIC
    scanned = img.transform((img_w, img_h), perspective_mode, coeffs, resample=resample)
    if enhance:
        scanned = _enhance_scan_color(scanned)

    writer = PdfWriter()
    for idx, src_page in enumerate(reader.pages):
        if idx != page_idx:
            writer.add_page(src_page)
            continue
        width = float(src_page.mediabox.width)
        height = float(src_page.mediabox.height)
        img_buf = io.BytesIO()
        scanned.save(img_buf, format="JPEG", quality=90)
        img_buf.seek(0)

        page_pdf = io.BytesIO()
        c = canvas.Canvas(page_pdf, pagesize=(width, height))
        c.drawImage(ImageReader(img_buf), 0, 0, width=width, height=height)
        c.showPage()
        c.save()
        scanned_reader = PdfReader(io.BytesIO(page_pdf.getvalue()))
        writer.add_page(scanned_reader.pages[0])

    out = io.BytesIO()
    writer.write(out)
    return _pdf_response(out.getvalue(), "scanned.pdf")


@router.post("/delete-pages")
async def delete_pages(file: UploadFile = File(...), pages: str = Form(...)):
    raw = await file.read()
    _check_size(raw)
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc
    total = len(reader.pages)
    keep = _parse_keep(pages, total)
    if not keep:
        raise HTTPException(400, "Sau khi xóa không còn trang nào")
    writer = PdfWriter()
    for idx in keep:
        writer.add_page(reader.pages[idx])
    out = io.BytesIO()
    writer.write(out)
    return _pdf_response(out.getvalue(), "edited.pdf")


@router.post("/merge")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(400, "Cần ít nhất 2 PDF")
    writer = PdfWriter()
    for f in files:
        raw = await f.read()
        _check_size(raw)
        try:
            reader = PdfReader(io.BytesIO(raw))
        except Exception as exc:
            raise HTTPException(400, f"PDF không hợp lệ: {f.filename}") from exc
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return _pdf_response(out.getvalue(), "merged.pdf")


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...), use_ocr: bool = Form(False)):
    raw = await file.read()
    _check_size(raw)
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        method = "pdf"
        if not text and use_ocr:
            png = _render_page_png(raw, i, scale=2.0)
            if png:
                ocr_text = await _ocr_remote(png, lang="vie+eng")
                if ocr_text:
                    text = ocr_text
                    method = "ocr"
        pages.append({"index": i + 1, "text": text, "method": method})
    return {"pages": pages, "count": len(pages)}


async def _translate_chunk(text: str, target: str) -> str:
    if not text.strip():
        return ""
    url = f"https://lingva.ml/api/v1/auto/{target}/{quote(text[:4500])}"
    try:
        async with aiohttp.ClientSession() as s:
            async with s.get(url, timeout=aiohttp.ClientTimeout(total=20)) as r:
                if r.status != 200:
                    return text
                data = await r.json()
                return data.get("translation") or text
    except Exception:
        return text


def _render_page_png(pdf_bytes: bytes, page_index: int, scale: float = 2.0) -> bytes | None:
    """Render 1 page PDF sang PNG (dùng cho OCR fallback)."""
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
        if page_index >= len(doc):
            return None
        page = doc[page_index]
        bmp = page.render(scale=scale)
        img = bmp.to_pil()
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        return buf.getvalue()
    except Exception:
        return None


def _average_corner_color(img: Image.Image, sample: int = 24) -> tuple[int, int, int]:
    """Ước lượng màu nền bằng trung bình 4 góc ảnh để fill vùng xoay."""
    rgb_img = img.convert("RGB")
    w, h = rgb_img.size
    sample = max(1, min(sample, w // 4 or 1, h // 4 or 1))
    boxes = [
        (0, 0, sample, sample),
        (w - sample, 0, w, sample),
        (0, h - sample, sample, h),
        (w - sample, h - sample, w, h),
    ]
    totals = [0, 0, 0]
    count = 0
    for box in boxes:
        for r, g, b in rgb_img.crop(box).getdata():
            totals[0] += r
            totals[1] += g
            totals[2] += b
            count += 1
    if not count:
        return (255, 255, 255)
    return tuple(int(v / count) for v in totals)


def _perspective_coeffs(dst_points: list[tuple[float, float]], src_points: list[tuple[float, float]]) -> list[float]:
    """Return PIL perspective coeffs mapping destination pixels to source pixels."""
    import numpy as np

    matrix = []
    values = []
    for (x, y), (u, v) in zip(dst_points, src_points):
        matrix.append([x, y, 1, 0, 0, 0, -u * x, -u * y])
        matrix.append([0, 0, 0, x, y, 1, -v * x, -v * y])
        values.extend([u, v])
    coeffs = np.linalg.solve(np.asarray(matrix, dtype=float), np.asarray(values, dtype=float))
    return coeffs.tolist()


def _enhance_scan_color(img: Image.Image) -> Image.Image:
    """Tăng cảm giác scan màu: cân sáng nhẹ, rõ chữ, giữ mộc đỏ."""
    out = ImageOps.autocontrast(img.convert("RGB"), cutoff=1)
    out = ImageEnhance.Contrast(out).enhance(1.08)
    out = ImageEnhance.Sharpness(out).enhance(1.12)
    out = ImageEnhance.Color(out).enhance(1.03)
    return out


async def _ocr_remote(img_bytes: bytes, lang: str = "vie+eng") -> str:
    """Gọi OCR microservice trên remote. Trả về text hoặc '' nếu lỗi."""
    import os
    ocr_url = os.environ.get("HAGENT_OCR_URL", "http://100.69.50.64:8011").rstrip("/")
    try:
        form = aiohttp.FormData()
        form.add_field("image", img_bytes, filename="page.png", content_type="image/png")
        form.add_field("lang", lang)
        async with aiohttp.ClientSession() as s:
            async with s.post(
                f"{ocr_url}/v1/ocr/extract",
                data=form,
                timeout=aiohttp.ClientTimeout(total=60),
            ) as r:
                if r.status != 200:
                    return ""
                data = await r.json()
                return (data.get("text") or "").strip()
    except Exception:
        return ""


@router.post("/translate")
async def translate_pdf(file: UploadFile = File(...), target_lang: str = Form("vi")):
    target = LANG_MAP.get(target_lang.lower())
    if not target:
        raise HTTPException(400, f"Ngôn ngữ không hỗ trợ: {target_lang}")
    raw = await file.read()
    _check_size(raw)
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(400, f"PDF không hợp lệ: {exc}") from exc
    chunks = []
    ocr_used_pages: list[int] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        ocr_used = False
        if not text.strip():
            png = _render_page_png(raw, i, scale=2.0)
            if png:
                text = await _ocr_remote(png, lang="vie+eng")
                ocr_used = bool(text.strip())
                if ocr_used:
                    ocr_used_pages.append(i + 1)
        if not text.strip():
            continue
        translated = await _translate_chunk(text, target)
        tag = " (OCR)" if ocr_used else ""
        chunks.append(f"--- Trang {i + 1}{tag} ---\n{translated}")
    if not chunks:
        raise HTTPException(400, "Không trích xuất được text để dịch (kể cả OCR)")
    out_name = (Path(file.filename or "doc").stem) + f".{target}.pdf"
    headers = {}
    if ocr_used_pages:
        headers["X-OCR-Pages"] = ",".join(str(p) for p in ocr_used_pages)
    return StreamingResponse(
        io.BytesIO(_build_text_pdf("\n\n".join(chunks), title=f"Bản dịch ({target_lang.upper()})")),
        media_type="application/pdf",
        headers={**headers, "Content-Disposition": f'attachment; filename="{quote(out_name)}"'},
    )


# === Translate job với SSE progress ===

async def _run_translate_job(
    job_id: str, raw: bytes, target: str, target_lang: str, filename: str,
    llm_settings: dict | None = None,
) -> None:
    """Chạy translation trong background, cập nhật job state để SSE stream ra ngoài.

    Pipeline: Đọc PDF → (OCR từng trang nếu rỗng) → [LLM correction nếu bật]
              → Smart chunking + dịch (LLM hoặc lingva.ml) → Build PDF.
    """
    job = _translate_jobs.get(job_id)
    if not job:
        return
    llm_settings = llm_settings or {"provider": "none"}
    llm_enabled = (llm_settings.get("provider") or "none").lower() not in ("none", "", "off")
    try:
        job["status"] = "reading"
        job["message"] = "Đang đọc file PDF..."
        job["step"] = 1
        job["step_label"] = "Đọc PDF"
        try:
            reader = PdfReader(io.BytesIO(raw))
        except Exception as exc:
            job["status"] = "error"
            job["error"] = f"PDF không hợp lệ: {exc}"
            return
        total = len(reader.pages)
        job["total_pages"] = total
        job["current_page"] = 0
        if total == 0:
            job["status"] = "error"
            job["error"] = "PDF rỗng (0 trang)"
            return

        # --- Stage 1+2: extract text, fallback OCR ---
        page_texts: list[tuple[int, str, bool]] = []  # (page_idx, text, ocr_used)
        for i, page in enumerate(reader.pages):
            job["current_page"] = i + 1
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            ocr_used = False
            if not text.strip():
                job["status"] = "ocr"
                job["step"] = 2
                job["step_label"] = f"OCR trang {i + 1}/{total}"
                job["message"] = f"Trang {i + 1}/{total} không có text, đang OCR..."
                png = _render_page_png(raw, i, scale=2.0)
                if png:
                    text = await _ocr_remote(png, lang="vie+eng")
                    ocr_used = bool(text.strip())
            if text.strip():
                page_texts.append((i, text, ocr_used))
            await asyncio.sleep(0)

        if not page_texts:
            job["status"] = "error"
            job["error"] = "Không trích xuất được text (kể cả OCR)"
            return

        # --- Stage 2.5 (optional): LLM clean OCR ---
        if llm_enabled:
            job["status"] = "cleaning"
            job["step"] = 3
            job["step_label"] = "Làm sạch OCR"
            job["message"] = "LLM đang sửa lỗi OCR (chính tả, dính chữ)..."
            cleaned: list[tuple[int, str, bool]] = []
            for idx, (pg, txt, ocr) in enumerate(page_texts):
                job["current_page"] = idx + 1
                job["message"] = f"LLM sửa trang {idx + 1}/{len(page_texts)}..."
                try:
                    new_txt = await _llm_correct_text(txt, llm_settings)
                except HTTPException as exc:
                    new_txt = txt
                    job["message"] = f"LLM lỗi, giữ nguyên OCR: {exc.detail}"
                cleaned.append((pg, new_txt, ocr))
                await asyncio.sleep(0)
            page_texts = cleaned

        # --- Stage 3: smart chunking + translate ---
        chunks: list[str] = []
        ocr_pages: list[int] = []
        for pg, text, ocr in page_texts:
            sub_chunks = _split_into_chunks(text, max_chars=2500)
            translated_parts: list[str] = []
            for sc in sub_chunks:
                base_step = 4 if llm_enabled else 3
                job["status"] = "translating"
                job["step"] = base_step
                job["step_label"] = f"Dịch trang {pg + 1}"
                job["message"] = f"Đang dịch trang {pg + 1} ({len(translated_parts) + 1}/{len(sub_chunks)} phần)..."
                translated_parts.append(
                    await _llm_translate_chunk(sc, target, llm_settings)
                )
                await asyncio.sleep(0)
            translated = "\n\n".join(translated_parts)
            if ocr:
                ocr_pages.append(pg + 1)
            tag = " (OCR)" if ocr else ""
            chunks.append(f"--- Trang {pg + 1}{tag} ---\n{translated}")

        # --- Stage 4: build PDF ---
        build_step = 5 if llm_enabled else 4
        job["status"] = "building"
        job["step"] = build_step
        job["step_label"] = "Tạo file PDF"
        job["message"] = "Đang tạo file PDF kết quả..."
        await asyncio.sleep(0)
        out_name = (Path(filename or "doc").stem) + f".{target_lang}.pdf"
        pdf_bytes = _build_text_pdf("\n\n".join(chunks), title=f"Bản dịch ({target_lang.upper()})")
        saved_path = _save_to_recent(pdf_bytes, out_name)

        job["status"] = "done"
        job["step"] = build_step + 1
        job["step_label"] = "Hoàn tất"
        job["message"] = "Hoàn tất!"
        job["filename"] = out_name
        job["saved_path"] = saved_path
        job["ocr_pages"] = ocr_pages
        job["llm_used"] = llm_enabled
        job["pdf_base64"] = base64.b64encode(pdf_bytes).decode()
    except Exception as exc:
        job["status"] = "error"
        job["error"] = f"Lỗi: {exc}"


@router.post("/translate-job")
async def translate_job_start(
    file: UploadFile = File(...),
    target_lang: str = Form("vi"),
    llm_settings_json: str | None = Form(None),
):
    """Khởi động job translate, trả về job_id để client mở SSE stream theo dõi.

    llm_settings_json: JSON string chứa LLM config (provider/model/api_key/base_url).
    """
    target = LANG_MAP.get(target_lang.lower())
    if not target:
        raise HTTPException(400, f"Ngôn ngữ không hỗ trợ: {target_lang}")
    raw = await file.read()
    _check_size(raw)
    llm_settings: dict = {"provider": "none"}
    if llm_settings_json:
        try:
            parsed = json.loads(llm_settings_json)
            if isinstance(parsed, dict):
                llm_settings = parsed
        except json.JSONDecodeError:
            raise HTTPException(400, "llm_settings_json không phải JSON hợp lệ")
    job_id = uuid.uuid4().hex[:12]
    _translate_jobs[job_id] = {
        "status": "queued",
        "step": 0,
        "step_label": "Chờ xử lý",
        "message": "Đang vào hàng đợi...",
        "current_page": 0,
        "total_pages": 0,
        "created": time.time(),
    }
    asyncio.create_task(
        _run_translate_job(
            job_id, raw, target, target_lang, file.filename or "doc.pdf",
            llm_settings=llm_settings,
        )
    )
    return {"job_id": job_id}


# === LLM correction endpoints ===

@router.post("/llm-correct")
async def llm_correct(request: Request):
    """Sửa text bằng LLM (test/debug). Body: {text, settings}."""
    body = await request.json()
    text = (body.get("text") or "").strip()
    settings = body.get("settings") or {"provider": "none"}
    if not text:
        raise HTTPException(400, "Thiếu text")
    corrected = await _llm_correct_text(text, settings)
    return {"original": text, "corrected": corrected, "length": len(corrected)}


@router.post("/test-llm")
async def test_llm(request: Request):
    """Test kết nối LLM. Body: {settings}."""
    body = await request.json()
    settings = body.get("settings") or {}
    provider = (settings.get("provider") or "none").lower()
    if provider in ("none", "", "off"):
        return {"ok": False, "error": "Provider = none (LLM tắt)"}
    model = settings.get("model") or "?"
    try:
        reply = await _llm_chat(
            [
                {"role": "system", "content": "Bạn là trợ lý. Trả lời NGẮN GỌN."},
                {"role": "user", "content": "Trả lời đúng 1 từ: OK"},
            ],
            settings,
            timeout=30,
        )
        return {"ok": True, "model": model, "reply": reply[:200]}
    except HTTPException as exc:
        return {"ok": False, "model": model, "error": f"{exc.status_code}: {exc.detail}"}


@router.get("/translate-job/{job_id}/stream")
async def translate_job_stream(job_id: str):
    """SSE stream tiến trình job. Events: progress, done, error, ping."""
    if EventSourceResponse is None:
        raise HTTPException(500, "Thiếu sse-starlette. Cài: pip install sse-starlette")
    if job_id not in _translate_jobs:
        raise HTTPException(404, "Job not found")

    async def event_gen():
        last_ping = time.time()
        try:
            while True:
                job = _translate_jobs.get(job_id)
                if not job:
                    yield {"event": "error", "data": json.dumps({"error": "Job expired"})}
                    break
                # Gửi event progress
                payload = {
                    "status": job["status"],
                    "step": job["step"],
                    "step_label": job["step_label"],
                    "message": job["message"],
                    "current_page": job["current_page"],
                    "total_pages": job["total_pages"],
                }
                yield {"event": "progress", "data": json.dumps(payload, ensure_ascii=False)}
                if job["status"] == "done":
                    done_payload = {
                        "filename": job["filename"],
                        "saved_path": job["saved_path"],
                        "ocr_pages": job.get("ocr_pages", []),
                        "pdf_base64": job["pdf_base64"],
                    }
                    yield {"event": "done", "data": json.dumps(done_payload, ensure_ascii=False)}
                    break
                if job["status"] == "error":
                    yield {"event": "error", "data": json.dumps({"error": job.get("error", "Unknown")}, ensure_ascii=False)}
                    break
                # Ping mỗi 15s để giữ connection
                if time.time() - last_ping > 15:
                    yield {"event": "ping", "data": "{}"}
                    last_ping = time.time()
                await asyncio.sleep(0.25)
        finally:
            # Cleanup job sau khi client đóng (giữ thêm 60s cho retry)
            async def _delayed_cleanup():
                await asyncio.sleep(60)
                _translate_jobs.pop(job_id, None)
            asyncio.create_task(_delayed_cleanup())

    return EventSourceResponse(event_gen())


@router.get("/translate-job/{job_id}")
async def translate_job_status(job_id: str):
    """Polling fallback nếu SSE không dùng được."""
    job = _translate_jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    return {
        "status": job["status"],
        "step": job["step"],
        "step_label": job["step_label"],
        "message": job["message"],
        "current_page": job["current_page"],
        "total_pages": job["total_pages"],
    }


@router.get("/recent")
async def recent_pdfs():
    """Scan common directories for recently modified PDF files."""
    home = Path.home()
    project_root = Path(__file__).resolve().parent.parent.parent.parent
    pdf_data_dir = project_root / "data" / "pdf"
    scan_dirs = [home / d for d in ("Documents", "Downloads", "Desktop", "")] + [pdf_data_dir]
    folder_labels = {
        "Documents": "Documents",
        "Downloads": "Downloads",
        "Desktop": "Desktop",
        "pdf": "PDF Editor",
    }
    results = []
    seen: set[str] = set()
    for d in scan_dirs:
        if not d.is_dir():
            continue
        try:
            for f in d.iterdir():
                if f.suffix.lower() == ".pdf" and f.is_file() and str(f) not in seen:
                    seen.add(str(f))
                    st = f.stat()
                    results.append({
                        "name": f.name,
                        "path": str(f),
                        "folder": folder_labels.get(f.parent.name, f.parent.name),
                        "size": st.st_size,
                        "modified": st.st_mtime,
                    })
        except PermissionError:
            continue
    results.sort(key=lambda x: x["modified"], reverse=True)
    return {"files": results[:30]}


@router.get("/open")
async def open_pdf_by_path(path: str):
    """Return PDF file bytes from the local filesystem."""
    p = Path(path).resolve()
    home = Path.home().resolve()
    if not str(p).startswith(str(home)):
        raise HTTPException(403, "Không được phép truy cập file ngoài thư mục home")
    if not p.exists() or p.suffix.lower() != ".pdf":
        raise HTTPException(404, "Không tìm thấy file PDF")
    return FileResponse(str(p), media_type="application/pdf",
                        headers={"Content-Disposition": f'attachment; filename="{quote(p.name)}"'})


@router.post("/save")
async def save_pdf(path: str = Form(...), file: UploadFile = File(...)):
    """Save PDF bytes to a local file path.

    If path is relative, save under the project's data/pdf/ directory.
    If absolute, path must be under user home.
    """
    home = Path.home().resolve()
    project_root = Path(__file__).resolve().parent.parent.parent.parent  # HAgent/
    pdf_data_dir = project_root / "data" / "pdf"
    candidate = Path(path)
    if candidate.is_absolute():
        p = candidate.resolve()
        if not str(p).startswith(str(home)):
            raise HTTPException(403, "Không được phép lưu file ngoài thư mục home")
    else:
        p = (pdf_data_dir / path).resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
    if p.suffix.lower() != ".pdf":
        p = p.with_suffix(".pdf")
    p.parent.mkdir(parents=True, exist_ok=True)
    content = await file.read()
    p.write_bytes(content)
    return {"path": str(p)}


def _inside_home(p: Path) -> bool:
    """Chỉ cho phép xoá/đổi tên file PDF nằm trong thư mục home user (an toàn)."""
    home = Path.home().resolve()
    try:
        p = p.resolve()
    except Exception:
        return False
    return str(p).startswith(str(home))


def _is_scan_dir(p: Path) -> bool:
    """True nếu file thuộc 1 trong các folder được /recent quét."""
    home = Path.home().resolve()
    scan_dirs = [home / d for d in ("Documents", "Downloads", "Desktop", "")] + [_pdf_data_dir().resolve()]
    try:
        p = p.resolve()
    except Exception:
        return False
    return any(str(p).startswith(str(d) + "/") or str(p) == str(d) for d in scan_dirs)


@router.post("/delete")
async def delete_pdf(path: str = Form(...)):
    """Xoá 1 file PDF trong các folder scan: Documents, Downloads, Desktop, Home, data/pdf/."""
    p = Path(path)
    if not _is_scan_dir(p):
        raise HTTPException(403, "Chỉ được xoá file PDF trong Documents/Downloads/Desktop/Home/PDF Editor")
    if not _inside_home(p):
        raise HTTPException(403, "Chỉ được xoá file trong thư mục home")
    if not p.exists() or p.suffix.lower() != ".pdf":
        raise HTTPException(404, "Không tìm thấy file PDF")
    try:
        p.unlink()
    except Exception as exc:
        raise HTTPException(500, f"Không xoá được: {exc}") from exc
    return {"ok": True, "path": str(p.resolve())}


@router.post("/rename")
async def rename_pdf(path: str = Form(...), new_name: str = Form(...)):
    """Đổi tên file PDF trong các folder scan: Documents, Downloads, Desktop, Home, data/pdf/."""
    p = Path(path)
    if not _is_scan_dir(p):
        raise HTTPException(403, "Chỉ được đổi tên file PDF trong Documents/Downloads/Desktop/Home/PDF Editor")
    if not _inside_home(p):
        raise HTTPException(403, "Chỉ được đổi tên file trong thư mục home")
    if not p.exists() or p.suffix.lower() != ".pdf":
        raise HTTPException(404, "Không tìm thấy file PDF")
    safe = re.sub(r"[^\w.\- ]+", "_", new_name).strip()
    if not safe:
        raise HTTPException(400, "Tên file rỗng")
    if not safe.lower().endswith(".pdf"):
        safe += ".pdf"
    new_p = p.with_name(safe)
    if new_p == p.resolve():
        return {"ok": True, "path": str(p.resolve()), "unchanged": True}
    if new_p.exists():
        raise HTTPException(409, f"File '{safe}' đã tồn tại")
    try:
        p.rename(new_p)
    except Exception as exc:
        raise HTTPException(500, f"Không đổi tên được: {exc}") from exc
    return {"ok": True, "old_path": str(p.resolve()), "path": str(new_p.resolve())}


@router.post("/compress")
async def compress_pdf(file: UploadFile = File(...)):
    """Nén PDF: FlateDecode content streams + tối ưu object stream. Trả về PDF mới."""
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "File rỗng")
    if raw[:4] != b"%PDF":
        raise HTTPException(400, "Không phải file PDF hợp lệ")
    original_size = len(raw)
    try:
        reader = PdfReader(io.BytesIO(raw))
        writer = PdfWriter()
        for page in reader.pages:
            try:
                page.compress_content_streams()
            except Exception:
                pass
            writer.add_page(page)
        out = io.BytesIO()
        try:
            writer.compress_identical_objects()
        except TypeError:
            pass
        writer.write(out)
        new_buf = out.getvalue()
    except Exception as exc:
        raise HTTPException(500, f"Lỗi nén PDF: {exc}") from exc

    saved_path = _save_to_recent(new_buf, "compressed.pdf")
    return StreamingResponse(
        io.BytesIO(new_buf),
        media_type="application/pdf",
        headers={
            "Content-Disposition": 'attachment; filename="compressed.pdf"',
            "X-Saved-Path": saved_path,
            "X-Original-Size": str(original_size),
            "X-Compressed-Size": str(len(new_buf)),
            "X-Saved-Bytes": str(max(0, original_size - len(new_buf))),
        },
    )


def _parse_keep(pages: str, total: int) -> list[int]:
    """Parse '2,5-7' as pages to DELETE; return 0-based indexes to keep."""
    drop: set[int] = set()
    for part in pages.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                a, b = [int(x) for x in part.split("-", 1)]
            except ValueError:
                continue
            for n in range(min(a, b), max(a, b) + 1):
                drop.add(n - 1)
        else:
            try:
                drop.add(int(part) - 1)
            except ValueError:
                continue
    return [i for i in range(total) if i not in drop]
