from fastapi import APIRouter, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse
import google.generativeai as genai
import fitz  # PyMuPDF
import os, uuid, threading, time, random
import requests
from urllib.parse import quote
from io import BytesIO
from datetime import datetime
from docx import Document
from pathlib import Path

router = APIRouter(prefix="/api/pdf", tags=["pdf-translate"])

# ==========================================================
# TASK MANAGEMENT
# ==========================================================

tasks = {}

def update_task(task_id, status, progress=0, result=None, error=None):
    tasks[task_id] = {
        "status": status,
        "progress": progress,
        "result": result,
        "error": error,
        "updated_at": datetime.now().isoformat()
    }

# ==========================================================
# API KEY ROTATION
# ==========================================================

api_keys = []
api_index = 0

def set_api_keys(keys):
    global api_keys
    api_keys = [k.strip() for k in keys.split(",") if k.strip()]
    print(f"[API] Loaded {len(api_keys)} API keys")

def get_api_key():
    global api_index
    if not api_keys:
        return os.getenv("GEMINI_API_KEY")
    key = api_keys[api_index]
    api_index = (api_index + 1) % len(api_keys)
    return key

# ==========================================================
# GEMINI TRANSLATION
# ==========================================================

def translate_text(text, src="auto", tgt="vi", ctx=""):
    api_key = get_api_key()
    if not api_key:
        print("[TRANSLATE] Error: No Gemini API Key configured.")
        return text

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("models/gemini-2.0-flash")

    prompt = f"""
Translate the following text from {src} to {tgt}.
Context: {ctx}
Output ONLY translated text.

{text}
"""

    for attempt in range(3):
        try:
            res = model.generate_content(prompt)
            if res and res.text:
                return res.text.strip()
            return text
        except Exception as e:
            print(f"[TRANSLATE] Lỗi: {e} — retry {attempt+1}/3")
            time.sleep(random.uniform(0.5, 1.5))

    return text

# ==========================================================
# LINGVA BASIC TRANSLATION
# ==========================================================

def translate_text_basic(text, tgt="vi"):
    if not text.strip():
        return ""
    url = f"https://lingva.ml/api/v1/auto/{tgt}/{quote(text[:4500])}"
    try:
        r = requests.get(url, timeout=15)
        if r.status_code == 200:
            data = r.json()
            return data.get("translation") or text
    except Exception as e:
        print(f"[BASIC TRANSLATE] Lỗi: {e}")
    return text

# ==========================================================
# PDF → TEXT EXTRACTION (USING PyMuPDF)
# ==========================================================

def pdf_extract_text(pdf_bytes):
    """
    Extract text from PDF using PyMuPDF (fitz).
    Works with encrypted PDFs and handles errors gracefully.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages = []
        
        for i in range(len(doc)):
            page = doc[i]
            txt = page.get_text().strip()
            pages.append({"page": i + 1, "text": txt})
        
        doc.close()
        return pages
    
    except Exception as e:
        print(f"[PDF EXTRACT] Lỗi: {e}")
        raise HTTPException(400, f"Không thể đọc PDF: {str(e)}. File có thể bị mã hóa, hỏng hoặc không hợp lệ.")

# ==========================================================
# GET OUTPUT DIR
# ==========================================================

def get_output_dir():
    # HAgent project root / data / pdf / output
    project_root = Path(__file__).resolve().parents[3]
    d = project_root / "data" / "pdf" / "output"
    d.mkdir(parents=True, exist_ok=True)
    return d

# ==========================================================
# SAVE DOCX
# ==========================================================

def save_docx(pages, filename, src, tgt):
    doc = Document()
    for p in pages:
        doc.add_heading(f"Page {p['page']}", level=1)
        doc.add_paragraph(p["text"])
        doc.add_page_break()

    out_dir = get_output_dir()
    out_name = f"translated_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}.docx"
    out_path = out_dir / out_name
    doc.save(str(out_path))

    return out_name, out_path

# ==========================================================
# AUTO EXPAND TEXTBOX FOR PDF
# ==========================================================

def draw_translated_text(page, rect, translated, base_font=11):
    """
    Draw translated text with fixed small font size for consistency.
    """
    
    # Load Vietnamese-supporting font
    font_name = "helv"
    dejavu_path = str(Path(__file__).resolve().parents[2] / "assets" / "fonts" / "DejaVuSans.ttf")
    font_paths = [
        ("dejavu", dejavu_path),
        ("arial", "C:\\Windows\\Fonts\\arial.ttf"),
        ("times", "C:\\Windows\\Fonts\\times.ttf"),
        ("tahoma", "C:\\Windows\\Fonts\\tahoma.ttf"),
        ("arial", "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"),
        ("noto", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
        ("dejavu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    
    for name, path in font_paths:
        if os.path.exists(path):
            try:
                page.insert_font(fontname=name, fontfile=path)
                font_name = name
                break
            except:
                pass
    
    # Erase original area (no border)
    page.draw_rect(rect, color=None, fill=(1, 1, 1), width=0)
    
    # Fixed small font size for all blocks (consistent)
    font_size = 6.5
    
    page.insert_textbox(
        rect,
        translated,
        fontsize=font_size,
        fontname=font_name,
        color=(0, 0, 0),
        align=0
    )
    
    return rect

# ==========================================================
# PDF DIRECT TRANSLATION – WITH LOG + AUTO BLOCK EXPAND
# ==========================================================

def process_pdf_task(task_id, pdf_bytes, filename, src, tgt, ctx, engine):
    try:
        update_task(task_id, "processing", 5)
        print(f"[TASK {task_id}] Bắt đầu dịch PDF: {filename} ({engine})")

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        print(f"[TASK {task_id}] Tổng trang: {total_pages}")

        # Count all lines for progress
        total_lines = sum(
            len(b.get("lines", []))
            for p in doc
            for b in p.get_text("dict")["blocks"]
            if b.get("type") == 0
        )
        if total_lines == 0:
            total_lines = 1 # avoid division by zero
        done = 0

        # PROCESS EACH PAGE
        for p_index, page in enumerate(doc):
            print(f"\n=== Trang {p_index+1}/{total_pages} ===")

            blocks = page.get_text("dict")["blocks"]
            text_blocks = [b for b in blocks if b.get("type") == 0]
            text_blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))

            print(f"[TASK {task_id}] Số block: {len(text_blocks)}")

            for b_index, block in enumerate(text_blocks, 1):
                # Extract text
                block_text = " ".join(
                    " ".join(span["text"] for span in line["spans"])
                    for line in block["lines"]
                ).strip()

                if not block_text:
                    continue

                preview = block_text[:80]
                print(f"[TASK {task_id}] Block {b_index}/{len(text_blocks)}: {preview}...")

                # Translate
                if engine == "basic":
                    translated = translate_text_basic(block_text, tgt)
                else:
                    translated = translate_text(block_text, src, tgt, ctx)
                preview_tr = translated[:80]
                print(f"[TASK {task_id}] → Dịch: {preview_tr}...")

                # RECT OF BLOCK
                x0 = min(line["bbox"][0] for line in block["lines"])
                y0 = min(line["bbox"][1] for line in block["lines"])
                x1 = max(line["bbox"][2] for line in block["lines"])
                y1 = max(line["bbox"][3] for line in block["lines"])
                rect = fitz.Rect(x0, y0, x1, y1)

                # Draw translated text with auto expansion
                draw_translated_text(page, rect, translated)

                # Progress update
                done += len(block.get("lines", []))
                pct = 5 + int(done / total_lines * 80)
                update_task(task_id, "processing", min(pct, 90))
                print(f"[TASK {task_id}] Progress: {pct}%")

        # SAVE OUTPUT
        update_task(task_id, "processing", 95)
        out_dir = get_output_dir()

        out_name = f"translated_direct_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
        out_path = out_dir / out_name

        doc.save(str(out_path))
        doc.close()

        print(f"[TASK {task_id}] Hoàn tất → {out_path}")

        update_task(task_id, "completed", 100, {
            "filename": out_name,
            "download_url": f"/download-pdf/{out_name}"
        })

    except Exception as e:
        print(f"[TASK {task_id}] LỖI: {e}")
        update_task(task_id, "failed", 0, error=str(e))

# ==========================================================
# ROUTES
# ==========================================================

@router.post("/translate-doc")
async def translate_doc(file: UploadFile = File(...),
                        source_lang: str = "auto",
                        target_lang: str = "vi",
                        api_keys: str = None,
                        context: str = "",
                        engine: str = "gemini"):
    """
    Translate PDF to DOCX format.
    Extracts text from PDF and translates it to a Word document.
    """
    if not file.filename.endswith(".pdf"):
        raise HTTPException(400, "Only PDF files are supported")

    if api_keys and engine == "gemini":
        set_api_keys(api_keys)

    try:
        # Read PDF content
        pdf_content = await file.read()
        
        # Extract text from PDF
        pages = pdf_extract_text(BytesIO(pdf_content))
        
        if not pages:
            raise HTTPException(400, "PDF không có nội dung văn bản hoặc file bị hỏng")
        
        # Translate each page
        translated = []
        for p in pages:
            if p["text"]:  # Only translate if there's text
                if engine == "basic":
                    translated_text = translate_text_basic(p["text"], target_lang)
                else:
                    translated_text = translate_text(p["text"], source_lang, target_lang, context)
            else:
                translated_text = ""
            
            translated.append({"page": p["page"], "text": translated_text})
        
        # Save to DOCX
        out_name, _ = save_docx(translated, file.filename.replace(".pdf",""), source_lang, target_lang)

        return {
            "success": True,
            "filename": out_name,
            "download_url": f"/download/{out_name}",
            "pages_translated": len(pages)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] translate-doc: {e}")
        raise HTTPException(500, f"Lỗi xử lý file: {str(e)}")


@router.post("/translate-pdf-direct")
async def translate_pdf_direct(
    file: UploadFile = File(...),
    source_lang: str = "auto",
    target_lang: str = "vi",
    api_keys: str = None,
    context: str = "",
    engine: str = "gemini"
):
    """
    Translate PDF directly (preserving layout).
    Returns task_id for async processing.
    """
    if not file.filename.endswith(".pdf"):
        raise HTTPException(400, "Only PDF files are supported")

    if api_keys and engine == "gemini":
        set_api_keys(api_keys)

    task_id = str(uuid.uuid4())
    update_task(task_id, "pending", 0)

    pdf_bytes = await file.read()

    threading.Thread(
        target=process_pdf_task,
        args=(task_id, pdf_bytes, file.filename, source_lang, target_lang, context, engine),
        daemon=True
    ).start()

    return {"success": True, "task_id": task_id}


@router.get("/task-status/{task_id}")
async def get_task_status(task_id: str):
    """
    Check the status of a translation task.
    """
    if task_id not in tasks:
        raise HTTPException(404, "Task not found")
    return tasks[task_id]


@router.get("/download/{filename}")
async def download_doc(filename: str):
    """
    Download translated DOCX file.
    """
    if ".." in filename or filename.startswith("/") or filename.startswith("\\"):
        raise HTTPException(400, "Invalid filename")
    out_dir = get_output_dir()
    path = out_dir / filename
    if not path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(str(path), filename=filename)


@router.get("/download-pdf/{filename}")
async def download_pdf(filename: str):
    """
    Download translated PDF file.
    """
    if ".." in filename or filename.startswith("/") or filename.startswith("\\"):
        raise HTTPException(400, "Invalid filename")
    out_dir = get_output_dir()
    path = out_dir / filename
    if not path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(str(path), media_type="application/pdf", filename=filename)
